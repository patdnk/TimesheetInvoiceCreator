__author__ = 'patdnk'

import random
import sys
import os
import datetime
from tkinter import *
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import askdirectory
from tkinter.messagebox import showerror
# from decimal import Decimal
from docx import Document
from dateutil import parser
from decimal import Decimal
from os.path import abspath, expanduser

#todo Add validation before creating timesheet/invoice
#todo Add database history management (read, recreate, reprint)


class Application:
    def __init__(self, parent):

        self.myParent = parent

        # define StringVars
        self.timesheetTemplateFilenameString = StringVar()
        self.timesheetTemplateLabelString = StringVar()
        self.timesheetDestinationDirectoryString = StringVar()
        self.timesheetDestinationLabelString = StringVar()
        self.timesheetDestinationPathString = StringVar()
        self.dateSelected = StringVar()
        self.invoiceTemplateFilenameString = StringVar()
        self.invoiceTemplateLabelString = StringVar()
        self.invoiceDestinationDirectoryString = StringVar()
        self.invoiceDestinationLabelString = StringVar()
        self.invoiceDestinationPathString = StringVar()

        # define Main frame
        self.mainFrame = Frame(parent)
        self.mainFrame.grid()

        rowNumber = 0

        # define Timesheet label
        self.timesheetLabel = Label(self.mainFrame, text="Timesheet", font="System 14 bold")
        self.timesheetLabel.grid(row=rowNumber, column=0)

        # define Invoice label
        self.invoiceLabel = Label(self.mainFrame, text="Invoice", font="System 14 bold")
        self.invoiceLabel.grid(row=rowNumber, column=1)

        rowNumber += 1

        # define Timesheet frame
        self.timesheetFrame = Frame(self.mainFrame, bg="white", borderwidth=1, relief=SUNKEN, width=310, height=550)
        self.timesheetFrame.grid(row=rowNumber, column=0, padx=5, pady=5)
        self.timesheetFrame.grid_propagate(0)

        # create timesheet widgets
        self.createTimesheetWidgets()

        # define Invoice frame
        self.invoiceFrame = Frame(self.mainFrame, bg="white", borderwidth=1, relief=SUNKEN, width=310, height=550)
        self.invoiceFrame.grid(row=rowNumber, column=1, padx=5, pady=5)
        self.invoiceFrame.grid_propagate(0)

        # create invoice widgets
        self.createInvoiceFrameWidgets()

        # create helper tuples
        self.normalHoursEntryList = [self.entryMondayNH, self.entryTuesdayNH, self.entryWednesdayNH, self.entryThursdayNH, self.entryFridayNH, self.entrySaturdayNH, self.entrySundayNH]
        self.overtimeHoursEntryList = [self.entryMondayOH, self.entryTuesdayOH, self.entryWednesdayOH, self.entryThursdayOH, self.entryFridayOH, self.entrySaturdayOH, self.entrySundayOH]

        self.daysTuple = [("nh_monday",self.entryMondayNH),("nh_tuesday",self.entryTuesdayNH),("nh_wednesday",self.entryWednesdayNH),("nh_thursday",self.entryThursdayNH),("nh_friday",self.entryFridayNH),("nh_saturday",self.entrySaturdayNH),("nh_sunday",self.entrySundayNH),("oh_monday",self.entryMondayOH),("oh_tuesday",self.entryTuesdayOH),("oh_wednesday",self.entryTuesdayOH),("oh_thursday",self.entryThursdayOH),("oh_friday",self.entryFridayOH),("oh_saturday",self.entrySaturdayOH),("oh_sunday",self.entrySundayOH)];
        self.daysPlaceholders = ["monday_date", "tuesday_date", "wednesday_date", "thursday_date", "friday_date", "saturday_date", "sunday_date"]

        self.invoiceEntriesTuple = [("invoice_number",self.entryInvoiceNumber),("invoice_description", self.entryInvoiceDescription),("unit(s)",self.entryUnitsTotal),("unit_price",self.entryUnitPrice)]


    def createTimesheetWidgets(self):

        # template frame
        self.timesheetTemplateFrame = Frame(self.timesheetFrame, bg="white", width=300)
        self.timesheetTemplateFrame.grid(padx=5, pady=5)

        templateRowNumber = 0

        self.timesheetTemplateFrameTitle = Label(self.timesheetTemplateFrame, text="Template", font="System 11", fg="darkgray")
        self.timesheetTemplateFrameTitle.grid(row=templateRowNumber, sticky=W, columnspan=2)
        templateRowNumber += 1

        self.selectTimesheetTemplateButton = Button(self.timesheetTemplateFrame, text = "Select", command=self.selectTimesheetTemplate)
        self.selectTimesheetTemplateButton.grid(row=templateRowNumber)

        # templateRowNumber += 1

        # self.entryTemplatePath = Entry(self.timesheetTemplateFrame, width=25)
        # self.entryTemplatePath.grid(row=templateRowNumber, column=1)

        self.templatePathLabel = Label(self.timesheetTemplateFrame, textvariable=self.timesheetTemplateLabelString, width=32, font="System 11", bg="white smoke")
        self.templatePathLabel.grid(row=templateRowNumber, column=1)


        # details frame
        self.timesheetDetailsFrame = Frame(self.timesheetFrame, bg="white", width=300)
        self.timesheetDetailsFrame.grid(padx=5, pady=5)

        detailRowNumber = 0

        self.timesheetDetailsFrameTitle = Label(self.timesheetDetailsFrame, text="Details", font="System 11", fg="darkgray")
        self.timesheetDetailsFrameTitle.grid(row=detailRowNumber, sticky=W)
        detailRowNumber += 1

        self.consultantName = Entry(self.timesheetDetailsFrame, width=35)
        self.consultantName.insert(0, 'Your name')
        self.consultantName.grid(row=detailRowNumber, sticky=W)
        detailRowNumber += 1

        self.clientName = Entry(self.timesheetDetailsFrame, width=35)
        self.clientName.insert(0, 'Client name')
        self.clientName.grid(row=detailRowNumber, sticky=W)
        detailRowNumber += 1

        # date picker frame
        self.dateFrame = Frame(self.timesheetFrame, bg="white", width=300)
        self.dateFrame.grid(padx=5, pady=5)

        dateRowNumber = 0

        self.timesheetDetailsFrameTitle = Label(self.dateFrame, text="Date", font="System 11", fg="darkgray")
        self.timesheetDetailsFrameTitle.grid(row=dateRowNumber, sticky=W)

        spaceFiller = Frame(self.dateFrame, width=258)
        spaceFiller.grid(row=dateRowNumber, column=1)

        dateRowNumber += 1

        # set placeholder text
        self.dateSelected.set("Select monday")

        self.dateMenu = OptionMenu(self.dateFrame, self.dateSelected, *self.getAllMondays(2015))
        # way to feed list/tuple to option menu
        # self.dateMenu = apply(OptionMenu, (self.dateFrame, self.dateSelected) + tuple(self.getAllMondays(2015)))
        self.dateMenu.grid(row=dateRowNumber, columnspan=2, sticky=W)

        # hours frame
        self.hoursFrame = Frame(self.timesheetFrame, bg="white", width=300)
        self.hoursFrame.grid(padx=5, pady=5, sticky=W)

        hoursRowNumber = 0

        self.timesheetDetailsFrameTitle = Label(self.hoursFrame, text="Hours", font="System 11", fg="darkgray")
        self.timesheetDetailsFrameTitle.grid(row=hoursRowNumber, columnspan=3, sticky=W)

        hoursRowNumber += 1

        #header labels
        self.labelDays = Label(self.hoursFrame, text="Day", font="System 12 bold")
        self.labelDays.grid(row=hoursRowNumber)
        self.labelNormalHours = Label(self.hoursFrame, text="Normal\nhours", font="System 12 bold")
        self.labelNormalHours.grid(row=hoursRowNumber, column=1)
        self.labelOvertimeHours = Label(self.hoursFrame, text="Overtime\nhours", font="System 12 bold")
        self.labelOvertimeHours.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        #labels and inputs
        self.labelMonday = Label(self.hoursFrame, text="Monday", font="System 12")
        self.labelMonday.grid(row=hoursRowNumber, sticky=E)
        self.entryMondayNH = Entry(self.hoursFrame, width=2)
        self.entryMondayNH.grid(row=hoursRowNumber, column=1)
        self.entryMondayOH = Entry(self.hoursFrame, width=2)
        self.entryMondayOH.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        self.labelTuesday = Label(self.hoursFrame, text="Tuesday", font="System 12")
        self.labelTuesday.grid(row=hoursRowNumber, sticky=E)
        self.entryTuesdayNH = Entry(self.hoursFrame, width=2)
        self.entryTuesdayNH.grid(row=hoursRowNumber, column=1)
        self.entryTuesdayOH = Entry(self.hoursFrame, width=2)
        self.entryTuesdayOH.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        self.labelWednesday = Label(self.hoursFrame, text="Wednesday", font="System 12")
        self.labelWednesday.grid(row=hoursRowNumber, sticky=E)
        self.entryWednesdayNH = Entry(self.hoursFrame, width=2)
        self.entryWednesdayNH.grid(row=hoursRowNumber, column=1)
        self.entryWednesdayOH = Entry(self.hoursFrame, width=2)
        self.entryWednesdayOH.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        self.labelThursday = Label(self.hoursFrame, text="Thursday", font="System 12")
        self.labelThursday.grid(row=hoursRowNumber, sticky=E)
        self.entryThursdayNH = Entry(self.hoursFrame, width=2)
        self.entryThursdayNH.grid(row=hoursRowNumber, column=1,)
        self.entryThursdayOH = Entry(self.hoursFrame, width=2)
        self.entryThursdayOH.grid(row=hoursRowNumber, column=2,)
        hoursRowNumber += 1

        self.labelFriday = Label(self.hoursFrame, text="Friday", font="System 12")
        self.labelFriday.grid(row=hoursRowNumber, sticky=E)
        self.entryFridayNH = Entry(self.hoursFrame,width=2)
        self.entryFridayNH.grid(row=hoursRowNumber, column=1)
        self.entryFridayOH = Entry(self.hoursFrame,width=2)
        self.entryFridayOH.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        self.labelSaturday = Label(self.hoursFrame, text="Saturday", font="System 12")
        self.labelSaturday.grid(row=hoursRowNumber, sticky=E)
        self.entrySaturdayNH = Entry(self.hoursFrame, width=2)
        self.entrySaturdayNH.grid(row=hoursRowNumber, column=1)
        self.entrySaturdayOH = Entry(self.hoursFrame, width=2)
        self.entrySaturdayOH.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        self.labelSunday = Label(self.hoursFrame, text="Sunday", font="System 12")
        self.labelSunday.grid(row=hoursRowNumber, sticky=E)
        self.entrySundayNH = Entry(self.hoursFrame, width=2)
        self.entrySundayNH.grid(row=hoursRowNumber, column=1)
        self.entrySundayOH = Entry(self.hoursFrame, width=2)
        self.entrySundayOH.grid(row=hoursRowNumber, column=2)
        hoursRowNumber += 1

        # destination frame
        self.timesheetSavePathFrame = Frame(self.timesheetFrame, bg="white", width=300)
        self.timesheetSavePathFrame.grid(padx=5, pady=5)

        templateRowNumber = 0

        self.timesheetDestinationFrameTitle = Label(self.timesheetSavePathFrame, text="Destination", font="System 11", fg="darkgray")
        self.timesheetDestinationFrameTitle.grid(row=templateRowNumber, sticky=W, columnspan=2)
        templateRowNumber += 1

        self.selectTimesheetDestinationButton = Button(self.timesheetSavePathFrame, text = "Select", command=self.selectTimesheetDestinationFolder)
        self.selectTimesheetDestinationButton.grid(row=templateRowNumber)

        # templateRowNumber += 1

        # self.entryTemplatePath = Entry(self.timesheetTemplateFrame, width=25)
        # self.entryTemplatePath.grid(row=templateRowNumber, column=1)

        self.templateDestinationPathLabel = Label(self.timesheetSavePathFrame, textvariable=self.timesheetDestinationLabelString, width=32, font="System 11", bg="white smoke")
        self.templateDestinationPathLabel.grid(row=templateRowNumber, column=1)

        #add the creation button
        self.createButton = Button(self.timesheetFrame, text = "Create", command=self.createTimesheet)
        self.createButton.grid(sticky=E)

    def createInvoiceFrameWidgets(self):

        # template frame
        self.invoiceTemplateFrame = Frame(self.invoiceFrame, bg="white", width=300)
        self.invoiceTemplateFrame.grid(padx=5, pady=5)

        templateRowNumber = 0

        self.invoiceTemplateFrameTitle = Label(self.invoiceTemplateFrame, text="Template", font="System 11", fg="darkgray")
        self.invoiceTemplateFrameTitle.grid(row=templateRowNumber, sticky=W, columnspan=2)
        templateRowNumber += 1

        self.selectInvoiceTemplateButton = Button(self.invoiceTemplateFrame, text="Select", command=self.selectInvoiceTemplate)
        self.selectInvoiceTemplateButton.grid(row=templateRowNumber)

        # templateRowNumber += 1

        # self.entryTemplatePath = Entry(self.timesheetTemplateFrame, width=25)
        # self.entryTemplatePath.grid(row=templateRowNumber, column=1)

        self.invoicePathLabel = Label(self.invoiceTemplateFrame, textvariable=self.invoiceTemplateLabelString, width=32, font="System 11", bg="white smoke")
        self.invoicePathLabel.grid(row=templateRowNumber, column=1)

        # invoice details frame
        self.invoiceDetailsFrame = Frame(self.invoiceFrame, bg="white", width=300)
        self.invoiceDetailsFrame.grid(padx=5, pady=5)

        detailRowNumber = 0

        self.invoiceDetailsFrameTitle = Label(self.invoiceDetailsFrame, text="Details", font="System 11", fg="darkgray")
        self.invoiceDetailsFrameTitle.grid(row=detailRowNumber, sticky=W)
        detailRowNumber += 1

        self.labelInvoiceNumber = Label(self.invoiceDetailsFrame, text="Invoice number", font="System 12")
        self.labelInvoiceNumber.grid(row=detailRowNumber, sticky=E)
        self.entryInvoiceNumber = Entry(self.invoiceDetailsFrame, width=5)
        self.entryInvoiceNumber.grid(row=detailRowNumber, column=1, sticky=W)
        detailRowNumber += 1

        self.labelInvoiceDescription = Label(self.invoiceDetailsFrame, text="Invoice description", font="System 12")
        self.labelInvoiceDescription.grid(row=detailRowNumber, sticky=E)
        self.entryInvoiceDescription = Entry(self.invoiceDetailsFrame, width=20)
        self.entryInvoiceDescription.grid(row=detailRowNumber, column=1, sticky=W)
        detailRowNumber += 1

        self.labelUnitsTotal = Label(self.invoiceDetailsFrame, text="Units total", font="System 12")
        self.labelUnitsTotal.grid(row=detailRowNumber, sticky=E)
        self.entryUnitsTotal = Entry(self.invoiceDetailsFrame, width=5)
        self.entryUnitsTotal.grid(row=detailRowNumber, column=1, sticky=W)
        detailRowNumber += 1

        self.labelUnitPrice = Label(self.invoiceDetailsFrame, text="Unit price", font="System 12")
        self.labelUnitPrice.grid(row=detailRowNumber, sticky=E)
        self.entryUnitPrice = Entry(self.invoiceDetailsFrame, width=5)
        self.entryUnitPrice.grid(row=detailRowNumber, column=1, sticky=W)
        detailRowNumber += 1



        # destination frame
        self.invoiceSavePathFrame = Frame(self.invoiceFrame, bg="white", width=300)
        self.invoiceSavePathFrame.grid(padx=5, pady=5)

        templateRowNumber = 0

        self.invoiceDestinationFrameTitle = Label(self.invoiceSavePathFrame, text="Destination", font="System 11", fg="darkgray")
        self.invoiceDestinationFrameTitle.grid(row=templateRowNumber, sticky=W, columnspan=2)
        templateRowNumber += 1

        self.selectInvoiceDestinationButton = Button(self.invoiceSavePathFrame, text = "Select", command=self.selectInvoiceDestinationFolder)
        self.selectInvoiceDestinationButton.grid(row=templateRowNumber)

        # templateRowNumber += 1

        # self.entryTemplatePath = Entry(self.timesheetTemplateFrame, width=25)
        # self.entryTemplatePath.grid(row=templateRowNumber, column=1)

        self.invoiceDestinationPathLabel = Label(self.invoiceSavePathFrame, textvariable=self.invoiceDestinationLabelString, width=32, font="System 11", bg="white smoke")
        self.invoiceDestinationPathLabel.grid(row=templateRowNumber, column=1)

        #add the creation button
        invoiceCreateButton = Button(self.invoiceFrame, text = "Create", command=self.createInvoice)
        invoiceCreateButton.grid(sticky=E)

    ###### Timesheet methods ######

    def getAllMondays(self, year):

        oneday = datetime.timedelta(days=1)
        oneweek = datetime.timedelta(days=7)

        start = datetime.date(year=year, month=1, day=1)
        while start.weekday() != 0:
            start += oneday

        self.mondays = []
        while start.year == year:
            self.mondays.append(start)
            start += oneweek

        print(self.mondays)
        return self.mondays

    def getAllHoursTotals(self):

        # normal hours total calculation with empty values check
        self.normalWeekHours = float(0)
        for entry in self.normalHoursEntryList:
            if entry.get():
                normalWeekHours += float(entry.get())

        # overtime hours total calculation with empty values check
        self.overtimeWeekHours = float(0)
        for entry in self.overtimeHoursEntryList:
            if entry.get():
                overtimeWeekHours += float(entry.get())

        print("Normal hours total:", normalWeekHours)
        print("Overtime hours total:", overtimeWeekHours)
        # print(overtimeWeekHours)
        # print(self.entryMondayNH.get())
        # print("Button pressed")

    def selectTimesheetTemplate(self):
        timesheetTemplateFilename = askopenfilename(filetypes=(("Word 2007 doc files", "*.docx"),
                                                               ("All files", "*.*")))
        if timesheetTemplateFilename:
            try:
                self.timesheetTemplateFilenameString.set(timesheetTemplateFilename)
                print(self.timesheetTemplateFilenameString)
                urlParts = timesheetTemplateFilename.rsplit("/", 2)
                # trimmedTemplatePath = ".../" + urlParts[1] + "/" + urlParts[2]
                trimmedTemplatePath = ".../" + urlParts[2]
                self.timesheetTemplateLabelString.set(trimmedTemplatePath)
                print(trimmedTemplatePath)
            except:
                showerror("Open Source File", "Failed to read file\n'%s'" % timesheetTemplateFilename)
            return

    def selectTimesheetDestinationFolder(self):
        timesheetSavingDirectory  = askdirectory(title="Please select a directory")
        if len(timesheetSavingDirectory) > 0:
            try:
                self.timesheetDestinationDirectoryString.set(timesheetSavingDirectory)
                urlParts = timesheetSavingDirectory.rsplit("/", 2)
                trimmedTemplatePath = ".../" + urlParts[2]
                self.timesheetDestinationLabelString.set(trimmedTemplatePath)
                print(timesheetSavingDirectory)
            except:
                showerror("Select directory", "Failed to open directory\n'%s'" % timesheetSavingDirectory)
            return

    def createTimesheet(self):
        # self.timesheetDocumentTemplatePath = "/users/patdynek/Documents/Maze Sys Ltd docs/templates/wa_timesheet_template.docx"
        # self.timesheetDocumentTemplatePath = abspath(expanduser("~/") + 'templates/wa_timesheet_template.docx')
        self.openTimesheetDocumentTemplate()
        self.insertHoursValues(self.daysTuple)
        self.insertDatesStartingFrom(self.dateSelected.get())
        self.insertTimesheetDetails()
        self.saveTimesheetDocument()

    #open timesheet document
    def openTimesheetDocumentTemplate(self):
        self.timesheetDocument = Document(self.timesheetTemplateFilenameString.get())

    #insert hours values
    def insertHoursValues(self,daysTuple):
        #insert hours from entries
        for table in self.timesheetDocument.tables:
            for cell in table._cells:

                for (phDay,entryDay) in daysTuple:
                    if phDay in cell.text:
                        for paragraph in cell.paragraphs:
                            print(paragraph.text)
                            if entryDay.get():
                                paragraph.text = entryDay.get()
                            else:
                                paragraph.text = "-"
        #sum up all hours
        totalNormalHours = 0
        totalOvertimeHours = 0
        for (phDay,entryDay) in daysTuple:
            if phDay[:2] == "nh":
                if entryDay.get():
                    totalNormalHours += int(entryDay.get())
            elif phDay[:2] == "oh":
                if entryDay.get():
                    totalOvertimeHours += int(entryDay.get())


        if (totalNormalHours >= 0 or totalOvertimeHours >= 0):
            for table in self.timesheetDocument.tables:
                for cell in table._cells:
                    if "nh_total" in cell.text:
                        for paragraph in cell.paragraphs:
                            if totalNormalHours > 0:
                                paragraph.text = str(totalNormalHours)
                            else:
                                paragraph.text = "-"
                    elif "oh_total" in cell.text:
                        for paragraph in cell.paragraphs:
                            if totalOvertimeHours > 0:
                                paragraph.text = str(totalOvertimeHours)
                            else:
                                paragraph.text = "-"

    def insertTimesheetDetails(self):
        for table in self.timesheetDocument.tables:
                for cell in table._cells:
                    if "consultant_name" in cell.text:
                        for paragraph in cell.paragraphs:
                            if self.consultantName.get():
                                paragraph.text = self.consultantName.get()
                    elif "client_name" in cell.text:
                        for paragraph in cell.paragraphs:
                            if self.clientName.get():
                                paragraph.text = self.clientName.get()


    def insertDatesStartingFrom(self,mondayDate):

        print(mondayDate)
        mondayDateDT = parser.parse(mondayDate)
        print(mondayDateDT)

        weekDays = []
        self.dates = []

        #get dates for next 7 days from monday
        for i in range(0,7):
            d = mondayDateDT + datetime.timedelta(days=i)
            weekDays.append(d)

        #create list with tuples (placeholder, date)
        self.dates = list(zip(self.daysPlaceholders, weekDays))

        if weekDays.__len__() == 7:
            for table in self.timesheetDocument.tables:
                for cell in table._cells:
                        for (phDay,dayDate) in self.dates:
                            if phDay in cell.text:
                                for paragraph in cell.paragraphs:
                                    print(paragraph.text)
                                    paragraph.text = dayDate.strftime("%d/%m/%Y")

    #save timesheet document
    def saveTimesheetDocument(self):
        weekStartDay = parser.parse(self.dateSelected.get())
        weekEndDay = weekStartDay + datetime.timedelta(days=7)

        nameAbbr = ""
        for word in self.consultantName.get().split():
            nameAbbr += word[0]

        timesheetName = "tsheet" + self.clientName.get().lower() + nameAbbr.lower() + "_" + weekStartDay.strftime("%d%m%y") + "-" + weekEndDay.strftime("%d%m%Y") # "tsheetcsrpd_060415-120415"
        documentSavePath = self.timesheetDestinationDirectoryString.get() + "/" + timesheetName + ".docx"
        # self.document.save("/users/patdynek/Documents/Maze Sys Ltd docs/templates/saved_timesheet.docx")
        self.timesheetDocument.save(documentSavePath)


    ###### Invoice methods ######

    def selectInvoiceTemplate(self):
        invoiceTemplateFilename = askopenfilename(filetypes=(("Word 2007 doc files", "*.docx"),
                                                               ("All files", "*.*")))
        if invoiceTemplateFilename:
            try:
                self.invoiceTemplateFilenameString.set(invoiceTemplateFilename)
                urlParts = invoiceTemplateFilename.rsplit("/", 2)
                # trimmedTemplatePath = ".../" + urlParts[1] + "/" + urlParts[2]
                trimmedTemplatePath = ".../" + urlParts[2]
                self.invoiceTemplateLabelString.set(trimmedTemplatePath)
                print(trimmedTemplatePath)
            except:
                showerror("Open Source File", "Failed to read file\n'%s'" % timesheetTemplateFilename)
            return

    def selectInvoiceDestinationFolder(self):
        timesheetSavingDirectory  = askdirectory(title="Please select a directory")
        if len(timesheetSavingDirectory) > 0:
            try:
                self.invoiceDestinationDirectoryString.set(timesheetSavingDirectory)
                urlParts = timesheetSavingDirectory.rsplit("/", 2)
                trimmedTemplatePath = ".../" + urlParts[2]
                self.invoiceDestinationLabelString.set(trimmedTemplatePath)
                print(timesheetSavingDirectory)
            except:
                showerror("Select directory", "Failed to open directory\n'%s'" % timesheetSavingDirectory)
            return

    #open invoice document
    def openInvoiceDocumentTemplate(self):
        self.invoiceDocument = Document(self.invoiceTemplateFilenameString.get())

    def createInvoice(self):
        self.openInvoiceDocumentTemplate()
        self.insertInvoiceDetails(self.invoiceEntriesTuple)
        self.saveInvoiceDocument()

    def saveInvoiceDocument(self):
        invoiceName = "Invoice_MSLTD_" + self.entryInvoiceNumber.get() + "_" + datetime.datetime.now().strftime("%d%m%y")
        documentSavePath = self.invoiceDestinationDirectoryString.get() + "/" + invoiceName + ".docx"
        self.invoiceDocument.save(documentSavePath)

    def insertInvoiceDetails(self,detailsTuple):

        weekStartDay = parser.parse(self.dateSelected.get())
        weekEndDay = weekStartDay + datetime.timedelta(days=7)

        unitsTotal = 0
        if self.entryUnitsTotal.get():
            # totalUnits = int(self.entryUnitsTotal.get())
            unitsTotal = int(self.entryUnitsTotal.get()) * int(self.entryUnitPrice.get())
            vatTotal = unitsTotal * 0.20
            grossTotal = unitsTotal * 1.20
            print("unitsTotal: " + str(unitsTotal) + " | vatTotal: " + str(vatTotal) + " | grossTotal: " + str(grossTotal))

        for table in self.invoiceDocument.tables:
            for cell in table._cells:

                for (phDay,invoiceDetailsEntry) in detailsTuple:
                    if phDay in cell.text:
                        for paragraph in cell.paragraphs:
                            print(paragraph.text)
                            if invoiceDetailsEntry.get():
                                paragraph.text = invoiceDetailsEntry.get()
                            else:
                                paragraph.text = "-"

                if "invoice_date" in cell.text:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text)
                        paragraph.text = datetime.datetime.now().strftime("%d %b %Y")

                if "invoice_period" in cell.text:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text)
                        paragraph.text = weekStartDay.strftime("%d %b %y") + "   TO   " + weekEndDay.strftime("%d %b %y")

                if "units_total" in cell.text:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text + " " + str(format(Decimal(unitsTotal),".2f")))
                        paragraph.text = str(format(Decimal(unitsTotal),".2f"))

                if "sub_total" in cell.text:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text + " " + str(format(Decimal(unitsTotal),".2f")))
                        paragraph.text = str(format(Decimal(unitsTotal),".2f"))

                if "vat_total" in cell.text:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text + " " + str(format(Decimal(vatTotal),".2f")))
                        paragraph.text = str(format(Decimal(vatTotal),".2f"))

                if "gross_total" in cell.text:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text + " " + str(format(Decimal(grossTotal),".2f")))
                        paragraph.text = str(format(Decimal(grossTotal),".2f"))



#####################
#### Application ####
#####################

root = Tk()
root.title("Timesheets & Invoices")
root.geometry("650x610")
root.config(padx=5, pady=5)
root.grid_propagate(0)
app = Application(root)
root.mainloop()